"""
研究成果导出 —— P3
====================================================================
把一次研究导成能交出去的东西：Markdown / HTML / JSON / Word。

**两条贯穿全文的约束**：

① **每一句摘录后面都必须挂着出处。** 导出物一旦离开这个软件，
   就没有"点开看原文"这个动作了 —— 那时候一句没有出处的话
   和一句编的话没有任何区别。所以导出格式里出处是**正文的一部分**，
   不是可以关掉的附录。

② **生成版和摘录版必须视觉上分得开。** 摘录是原文逐字，生成是模型改写。
   导出后混在一起，三个月后你自己都分不清哪句是谁说的。
   所以生成版一律带一个显式的段落标记。

**PDF 不在这里做**：生成排版正确、中文字体正确的 PDF 需要一整套排版引擎
（reportlab 要自己塞字体、weasyprint 要装 GTK）。而桌面端本来就是
Chromium —— 它的 `printToPDF` 排版质量比任何 Python PDF 库都好，
中文字体也是现成的。所以这里出 HTML，PDF 由桌面端打印。
**这不是偷懒，是把活交给已经擅长它的那一层。**
"""

from __future__ import annotations

import json
import re
from datetime import UTC, datetime
from typing import Any

FORMATS = ("markdown", "html", "json", "docx", "single-html")

#: E1 简报模板。四种排法，**同一批事实换个组织方式**，
#: 不重新提炼、不改任何一句摘录 —— 换模板永远不该改变内容
TEMPLATES = ("points", "timeline", "compare", "qa")


def export_research(
    payload: dict[str, Any],
    *,
    fmt: str = "markdown",
    title: str | None = None,
    include_excluded: bool = False,
    template: str = "points",
) -> tuple[str | bytes, str, str]:
    """
    返回 `(内容, 文件扩展名, MIME)`。

    `payload` 是 `/api/web/research` 的响应体（或项目里存的某次 run）。
    `template` 见 `TEMPLATES`（E1）；`single-html` 是 E6 的离线单文件。
    """
    f = fmt if fmt in FORMATS else "markdown"
    t = (title or payload.get("query") or "研究简报").strip()[:120]
    if f == "json":
        return (
            json.dumps(payload, ensure_ascii=False, indent=2),
            "json",
            "application/json; charset=utf-8",
        )
    md = _to_markdown(
        payload, title=t, include_excluded=include_excluded, template=template
    )
    if f == "markdown":
        return md, "md", "text/markdown; charset=utf-8"
    if f == "html":
        return _to_html(md, title=t), "html", "text/html; charset=utf-8"
    if f == "single-html":
        return _to_single_html(payload, md, title=t), "html", "text/html; charset=utf-8"
    return _to_docx(payload, title=t), "docx", (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ────────────────────────────────────────────────────────────────
# E1 四种简报模板
# ────────────────────────────────────────────────────────────────
def render_template(payload: dict[str, Any], template: str) -> list[str]:
    """
    按模板把简报重排成 Markdown 行。

    🔴 **四个模板用的是同一批摘录，一个字都不改**。换模板改的只是
    「先看什么后看什么」和「怎么分组」。如果换个模板结论就变了，
    那说明其中至少一个模板在偷偷做提炼 —— 那正是这个项目从头到尾
    在拒绝的事。
    """
    b = payload.get("briefing") or {}
    tpl = template if template in TEMPLATES else "points"
    lines: list[str] = []

    consensus = b.get("consensus") or []
    disputes = b.get("disputes") or []
    topics = b.get("topics") or []

    if tpl == "timeline":
        # 时间线式：把所有带日期的摘录按时间排，没日期的单独放最后
        dated: list[tuple[str, dict[str, Any]]] = []
        undated: list[dict[str, Any]] = []
        for grp in (consensus, disputes):
            for item in grp:
                for ev in (item.get("evidence") or [item]):
                    d = str(ev.get("published") or "")
                    (dated.append((d, ev)) if d else undated.append(ev))
        dated.sort(key=lambda x: x[0])
        lines.append("## 时间线")
        lines.append("")
        for d, ev in dated:
            lines.append(f"- **{d[:10]}** {ev.get('text') or ''} "
                         f"[{ev.get('site') or '来源'}]({ev.get('url') or ''})")
        if undated:
            lines += ["", "### 没有日期的", ""]
            for ev in undated[:20]:
                lines.append(f"- {ev.get('text') or ''} "
                             f"[{ev.get('site') or '来源'}]({ev.get('url') or ''})")
            lines += ["", "> 这些条目的来源页面没有给出发布时间，"
                          "**不是被隐藏了**，是原站就没写", ""]
        return lines

    if tpl == "compare":
        # 对比表式：一行一个说法，列出支持方和反对方
        lines += ["## 说法对照表", "",
                  "| 说法 | 这么说的 | 有异议的 |", "|---|---|---|"]
        for item in disputes:
            sides = item.get("sides") or item.get("evidence") or []
            claim = str(item.get("topic") or item.get("keyword") or "")[:40]
            a = "、".join(str(s.get("site") or "") for s in sides[:3])
            bside = "、".join(str(s.get("site") or "") for s in sides[3:6])
            lines.append(f"| {claim} | {a or '—'} | {bside or '—'} |")
        for item in consensus[:20]:
            claim = str(item.get("topic") or item.get("keyword") or "")[:40]
            sites = "、".join(
                str(e.get("site") or "") for e in (item.get("evidence") or [])[:3]
            )
            lines.append(f"| {claim} | {sites or '—'} | 没找到异议 |")
        lines += ["", "> 「没找到异议」**不等于没有异议** —— "
                      "只表示这一轮检索里没有出现反驳材料", ""]
        return lines

    if tpl == "qa":
        # 问答式：每个主题变成一个问句，摘录当答案
        lines += ["## 问答", ""]
        for t in topics or (consensus + disputes):
            kw = str(t.get("keyword") or t.get("topic") or "").strip()
            if not kw:
                continue
            lines += [f"### 关于「{kw}」，各方是怎么说的？", ""]
            for ev in (t.get("evidence") or [])[:5]:
                lines.append(f"- {ev.get('text') or ''} "
                             f"—— [{ev.get('site') or '来源'}]({ev.get('url') or ''})")
            lines.append("")
        return lines

    return []          # points = 走原来的默认排版


# ────────────────────────────────────────────────────────────────
# Markdown（其余格式都从它派生）
# ────────────────────────────────────────────────────────────────
def _tail_sections(payload: dict[str, Any], *, include_excluded: bool) -> list[str]:
    """
    核查 + 全部来源 + 已排除。**四个模板共用这一段** ——
    换模板换的是主体的组织方式，出处清单和核查结论一个都不能少。
    """
    lines: list[str] = []
    v = payload.get("verification") or {}
    if v:
        lines += ["", "## 核查", "",
                  f"档位：`{v.get('level')}`　{v.get('note') or ''}", ""]
        for c in v.get("claims") or []:
            con = c.get("controversy") or {}
            lines.append(
                f"- **{c.get('claim') or ''}** —— 支持 {len(c.get('support') or [])}"
                f" ／ 反驳 {len(c.get('refute') or [])}"
                + (f"　争议度 {con.get('score')}" if con.get("score") is not None else "")
            )
        lines.append("")

    lines += ["## 全部来源", ""]
    for i, c in enumerate(payload.get("results") or [], 1):
        tr = c.get("trust") or {}
        lines.append(
            f"{i}. [{c.get('title')}]({c.get('url')}) — {c.get('site')}"
            f"　`{tr.get('tierLabel') or '未收录'}`"
        )
    if include_excluded and payload.get("excluded"):
        lines += ["", "## 已排除（折叠掉的，附原因）", ""]
        for c in payload["excluded"]:
            tr = c.get("trust") or {}
            lines.append(
                f"- [{c.get('title')}]({c.get('url')}) — {c.get('site')}"
                f"：{'；'.join(tr.get('reasons') or [])}"
            )
    return lines


def _to_markdown(
    payload: dict[str, Any], *, title: str, include_excluded: bool,
    template: str = "points",
) -> str:
    b = payload.get("briefing") or {}
    v = payload.get("verification") or {}
    lines: list[str] = [
        f"# {title}",
        "",
        f"> 导出时间：{datetime.now(UTC).astimezone().strftime('%Y-%m-%d %H:%M')}　"
        f"检索用时 {int(payload.get('elapsedMs') or 0) / 1000:.1f}s　"
        f"抓取正文 {payload.get('fetched') or 0} 篇",
        "",
    ]

    # E1：非默认模板走另一套排版，然后**照样接上后面的核查和来源清单** ——
    # 换模板只该换主体的组织方式，不该把核查结果一起换没了
    if template != "points":
        alt = render_template(payload, template)
        if alt:
            lines += alt
            lines += _tail_sections(payload, include_excluded=include_excluded)
            return "\n".join(lines)

    # 检索过程（S5 多轮）—— 放最前面。读简报的人有权先知道这些结论是怎么来的
    rounds = payload.get("rounds") or []
    if len(rounds) > 1:
        lines += ["## 检索过程", ""]
        for r in rounds:
            qs = r.get("queries") or []
            lines.append(f"**第 {r.get('round')} 轮**（新增 {r.get('newResults', 0)} 条结果）")
            for q in qs:
                lines.append(f"- `{q.get('text')}` — {q.get('why') or ''}")
            if r.get("skipped"):
                lines.append(f"- （没有第二轮：{r['skipped']}）")
            lines.append("")

    # 可信度概览
    ts = payload.get("trustSummary") or {}
    if ts:
        lines += ["## 结果成色", "", f"{ts.get('note') or ''}", ""]
        by_tier = ts.get("byTier") or {}
        if by_tier:
            lines.append("| 来源级别 | 条数 |")
            lines.append("|---|---|")
            for k, n in by_tier.items():
                lines.append(f"| {k} | {n} |")
            lines.append("")

    # 核查（V 组）
    if v:
        lines += ["## 核查", "", f"档位：`{v.get('level')}`　{v.get('note') or ''}", ""]
        for c in v.get("counterEvidence") or []:
            lines.append(
                f"- ⚠️ 反面材料：[{c.get('title')}]({c.get('url')}) — {c.get('site')}"
            )
        ret = v.get("retracted") or {}
        for doi, info in ret.items():
            lines.append(f"- 🔴 **已撤稿**：{info.get('title')}（DOI {doi}）")
        origin = v.get("origin") or {}
        if origin.get("note"):
            lines += ["", f"**溯源**：{origin['note']}"]
            e = origin.get("earliest") or {}
            if e:
                lines.append(
                    f"最早可查：[{e.get('title')}]({e.get('url')})"
                    f"（{e.get('published')}，{e.get('site')}）"
                )
        for cl in v.get("claims") or []:
            lines += ["", f"**断言**：{cl.get('claim')}"]
            lines.append(
                f"- 结论：`{cl.get('verdict')}`　支持 {cl.get('supportCount')}／"
                f"反驳 {cl.get('refuteCount')}　{cl.get('note') or ''}"
            )
            for s in cl.get("refute") or []:
                lines.append(f"  - 反驳：[{s.get('title')}]({s.get('url')})")
        lines.append("")

    # 简报正文
    lines += ["## 共识（原文摘录）", ""]
    if not (b.get("consensus")):
        lines.append("_没有找到 2 个以上独立站点互相印证的说法。_")
    for c in b.get("consensus") or []:
        lines.append(f"### {c.get('topic')}（{c.get('independentSites')} 个独立站点）")
        for e in c.get("evidence") or []:
            lines.append(f"- 「{e.get('text')}」")
            lines.append(f"  — [{e.get('title') or e.get('site')}]({e.get('url')})"
                         f"　{e.get('tier') or ''}")
        lines.append("")

    if b.get("disputes"):
        lines += ["## 分歧（并排放，不下结论）", ""]
        for d in b["disputes"]:
            lines.append(f"### {d.get('topic')}")
            for pair in d.get("conflicts") or []:
                a, bb = pair.get("a") or {}, pair.get("b") or {}
                lines.append(f"- A：「{a.get('text')}」— [{a.get('site')}]({a.get('url')})")
                lines.append(f"- B：「{bb.get('text')}」— [{bb.get('site')}]({bb.get('url')})")
                lines.append("")

    matrix = b.get("matrix") or {}
    if matrix.get("sites"):
        lines += ["## 一致性矩阵", "", f"{matrix.get('note') or ''}", ""]
        lines.append("| 话题 | " + " | ".join(matrix["sites"]) + " |")
        lines.append("|---" * (len(matrix["sites"]) + 1) + "|")
        sym = {"positive": "✔", "negative": "✘", "mixed": "～", "silent": ""}
        for topic, row in zip(matrix.get("topics") or [], matrix.get("cells") or [],
                              strict=False):
            lines.append(
                f"| {topic} | " + " | ".join(sym.get(c.get("stance"), "") for c in row) + " |"
            )
        lines.append("")

    if b.get("numbers"):
        lines += ["## 关键数据", ""]
        for n in b["numbers"]:
            lines.append(f"- **{n.get('value')}** —「{n.get('sentence')}」"
                         f"（[{n.get('site')}]({n.get('url')})）")
        lines.append("")

    if b.get("timeline"):
        lines += ["## 时间线", ""]
        for t in b["timeline"]:
            lines.append(f"- {t.get('published')} — [{t.get('title')}]({t.get('url')})")
        lines.append("")

    if b.get("openQuestions"):
        lines += ["## 还没查清", ""]
        lines += [f"- {q}" for q in b["openQuestions"]]
        lines.append("")

    # 生成版简报：**显式隔离**
    gen = payload.get("generated") or payload.get("generatedBriefing")
    if gen:
        lines += [
            "## AI 生成版（不是原文，是模型改写的）", "",
            "> 下面这一段由模型根据上面的摘录写成。**它可能改变了原意**，"
            "有疑问一律以上面的原文摘录为准。", "",
            str(gen.get("text") if isinstance(gen, dict) else gen), "",
        ]

    # 来源清单
    lines += ["## 全部来源", ""]
    for i, c in enumerate(payload.get("results") or [], 1):
        tr = c.get("trust") or {}
        lines.append(
            f"{i}. [{c.get('title')}]({c.get('url')}) — {c.get('site')}"
            f"　`{tr.get('tierLabel') or '未收录'}`"
            + (f"　⚠️ {'、'.join(tr.get('farmFlags') or [])}" if tr.get("farmFlags") else "")
        )
    if include_excluded and payload.get("excluded"):
        lines += ["", "## 已排除（折叠掉的，附原因）", ""]
        for c in payload["excluded"]:
            tr = c.get("trust") or {}
            lines.append(
                f"- [{c.get('title')}]({c.get('url')}) — {c.get('site')}"
                f"：{'；'.join(tr.get('reasons') or [])}"
            )
    return "\n".join(lines).rstrip() + "\n"


# ────────────────────────────────────────────────────────────────
# HTML（桌面端据此打印 PDF）
# ────────────────────────────────────────────────────────────────
#: 字体和字号跟应用内保持一致（锚点 5）：西文 Times New Roman、
#: 汉字宋体、标题更大的宋体。导出物换一套字体会让它看起来不像
#: 同一个软件出来的东西
_CSS = """
:root { color-scheme: light; }
body { font: 16px/1.75 "Times New Roman", "SimSun", 宋体, serif;
       color: #1F2933; background: #FAF9F6;
       max-width: 52em; margin: 2.5em auto; padding: 0 1.5em; }
h1, h2, h3 { font-family: "Source Han Serif SC", "SimSun", 宋体, serif;
             color: #0F4C8C; line-height: 1.35; }
h1 { font-size: 29.33px; border-bottom: 2px solid #0F4C8C; padding-bottom: .3em; }
h2 { font-size: 21.33px; margin-top: 1.8em; }
h3 { font-size: 20px; }
blockquote { border-left: 3px solid #C8871B; margin: 1em 0; padding: .2em 1em;
             color: #52606D; background: #fff; }
a { color: #0F4C8C; }
code { background: #EEF1F4; padding: .1em .35em; border-radius: 3px;
       font-family: Consolas, Menlo, monospace; font-size: .9em; }
table { border-collapse: collapse; margin: 1em 0; width: 100%; font-size: 14px; }
th, td { border: 1px solid #CBD2D9; padding: .4em .6em; text-align: left; }
th { background: #EEF1F4; }
li { margin: .25em 0; }
@media print { body { background: #fff; margin: 0; max-width: none; } a { color: #000; } }
"""


def _to_html(md: str, *, title: str) -> str:
    """
    极简 Markdown → HTML。

    **刻意不引第三方 markdown 库**：这里要渲染的语法是我们自己在
    `_to_markdown` 里写出来的，一共就那么七八种。为一个已知的、
    封闭的子集背一个依赖不划算 —— 而且第三方库默认允许内联 HTML，
    那等于把抓回来的网页标题当 HTML 渲染，是个实打实的注入面。
    """
    out: list[str] = []
    in_ul = False
    in_table = False
    for raw in md.split("\n"):
        line = _esc(raw)
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue  # 分隔行
            if not in_table:
                out.append("<table>")
                in_table = True
                out.append("<tr>" + "".join(f"<th>{_inline(c)}</th>" for c in cells) + "</tr>")
                continue
            out.append("<tr>" + "".join(f"<td>{_inline(c)}</td>" for c in cells) + "</tr>")
            continue
        if in_table:
            out.append("</table>")
            in_table = False

        if line.startswith("- ") or re.match(r"^\d+\. ", line):
            if not in_ul:
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{_inline(re.sub(r'^(- |\d+\. )', '', line))}</li>")
            continue
        if in_ul:
            out.append("</ul>")
            in_ul = False

        if line.startswith("### "):
            out.append(f"<h3>{_inline(line[4:])}</h3>")
        elif line.startswith("## "):
            out.append(f"<h2>{_inline(line[3:])}</h2>")
        elif line.startswith("# "):
            out.append(f"<h1>{_inline(line[2:])}</h1>")
        elif line.startswith("> "):
            out.append(f"<blockquote>{_inline(line[2:])}</blockquote>")
        elif line.strip():
            out.append(f"<p>{_inline(line)}</p>")
    if in_ul:
        out.append("</ul>")
    if in_table:
        out.append("</table>")
    return (
        "<!doctype html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\">"
        f"<title>{_esc(title)}</title><style>{_CSS}</style></head><body>"
        + "\n".join(out)
        + "</body></html>"
    )


def _esc(s: str) -> str:
    return (
        s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )


_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _inline(s: str) -> str:
    """行内语法。链接的 href 只允许 http/https —— 别把 `javascript:` 放进去。"""

    def link(m: re.Match[str]) -> str:
        href = m.group(2)
        if not href.startswith(("http://", "https://")):
            return m.group(1)
        return f'<a href="{href}" target="_blank" rel="noopener noreferrer">{m.group(1)}</a>'

    s = _LINK.sub(link, s)
    s = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", s)
    s = re.sub(r"`([^`]+)`", r"<code>\1</code>", s)
    return s


# ────────────────────────────────────────────────────────────────
# Word
# ────────────────────────────────────────────────────────────────
def _to_docx(payload: dict[str, Any], *, title: str) -> bytes:
    """
    Word 导出。`python-docx` 没装时**明确抛出可读的错误**，
    而不是悄悄退回 Markdown —— 用户点的是"导出 Word"，
    给他一个 .md 文件是最糟的处理。
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:  # noqa: BLE001
        raise RuntimeError(
            "导出 Word 需要 python-docx，去分析中心的依赖面板装一下（约 250 KB）"
        ) from e

    b = payload.get("briefing") or {}
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(12)  # 小四

    doc.add_heading(title, level=0)
    doc.add_paragraph(
        f"导出时间 {datetime.now(UTC).astimezone().strftime('%Y-%m-%d %H:%M')}　"
        f"抓取正文 {payload.get('fetched') or 0} 篇"
    )

    def section(name: str) -> None:
        doc.add_heading(name, level=1)

    section("共识（原文摘录）")
    for c in b.get("consensus") or []:
        doc.add_heading(f"{c.get('topic')}（{c.get('independentSites')} 个独立站点）", level=2)
        for e in c.get("evidence") or []:
            doc.add_paragraph(f"「{e.get('text')}」", style="List Bullet")
            doc.add_paragraph(f"　出处：{e.get('title') or e.get('site')} — {e.get('url')}")

    if b.get("disputes"):
        section("分歧")
        for d in b["disputes"]:
            doc.add_heading(str(d.get("topic")), level=2)
            for pair in d.get("conflicts") or []:
                for side in ("a", "b"):
                    ev = pair.get(side) or {}
                    doc.add_paragraph(
                        f"{side.upper()}：「{ev.get('text')}」— {ev.get('url')}",
                        style="List Bullet",
                    )

    v = payload.get("verification") or {}
    if v:
        section("核查")
        doc.add_paragraph(str(v.get("note") or ""))
        for c in v.get("counterEvidence") or []:
            doc.add_paragraph(f"反面材料：{c.get('title')} — {c.get('url')}",
                              style="List Bullet")

    section("全部来源")
    for i, c in enumerate(payload.get("results") or [], 1):
        tr = c.get("trust") or {}
        doc.add_paragraph(
            f"{i}. {c.get('title')} — {c.get('site')}"
            f"（{tr.get('tierLabel') or '未收录'}）\n　{c.get('url')}"
        )

    import io

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_filename(title: str, ext: str) -> str:
    """
    文件名清洗。Windows 不允许 `\\/:*?"<>|`，而研究标题里这些字符很常见
    （「A vs B：怎么选？」）—— 不清洗的话保存会直接失败。
    """
    name = re.sub(r'[\\/:*?"<>|\x00-\x1f]', "_", title).strip(" .") or "研究简报"
    return f"{name[:80]}.{ext}"


# ────────────────────────────────────────────────────────────────
# E6 单文件 HTML（离线可开，证据全内嵌）
# ────────────────────────────────────────────────────────────────
#: 单文件版额外的样式：加一个可折叠的「原始证据」区，
#: 以及 E5 要的**可点引用锚点**
_SINGLE_EXTRA_CSS = """
details.syn-raw { margin: 2em 0 0; border-top: 1px solid #d8d2c4; padding-top: 1em; }
details.syn-raw summary { cursor: pointer; font-size: 14pt; }
.syn-src { margin: .6em 0; padding: .6em .8em; background: #faf8f3;
           border-left: 3px solid #c8871b; }
.syn-src a { word-break: break-all; }
.syn-meta { color: #6b6558; font-size: 10.5pt; }
:target { background: #fff6e0; outline: 2px solid #c8871b; }
/* 指向文件内部证据区的引用链接。**和外链长得不一样**——
   用户要能一眼看出"这个点了会跳到下面"而不是"点了会开浏览器" */
a.syn-cite { border-bottom: 1px dashed #c8871b; text-decoration: none; }
a.syn-cite::after { content: "↓"; font-size: 0.8em; vertical-align: super; color: #c8871b; }
/* E5 打印成 PDF 时的样式。**证据区必须展开**（默认 open 就是为这个）——
   `details` 折叠着打印出来会丢掉整个内嵌证据区，而那正是这份文件的价值所在。
   ⚠️ 上面那条 `display:none` 是早期版本留下的，已改成展开打印。 */
@media print {
  details.syn-raw { display: block; }
  details.syn-raw > summary { list-style: none; }
  a { color: #0f4c8c; text-decoration: underline; }
  /* 只给外链补印真实网址：PDF 里链接可点是好的，但**打印出来的纸不能点**，
     纸上只剩一个"点这里"就等于把来源弄丢了。内部锚点（#src-3）不补，那是噪音 */
  a[href^="http"]::after { content: " <" attr(href) ">"; font-size: 9pt; word-break: break-all; }
  /* 内部引用在纸上印出**它指向第几条证据**，而不是印一个跳不动的 `#src-3`。
     纸不能点，所以必须把"跳过去会看到什么"变成看得见的文字 */
  a.syn-cite::after { content: " (见证据 " attr(href) ")"; font-size: 9pt; }
  /* 屏幕上高亮跳转目标用的是底色，纸上换成左侧竖线 —— 底色打印出来是一片灰 */
  :target { background: transparent; outline: none; border-left: 3px solid #0f4c8c; padding-left: 6px; }
  .syn-src { break-inside: avoid; }
}
"""


def _to_single_html(payload: dict[str, Any], md: str, *, title: str) -> str:
    """
    E6 —— 一个文件、双击就能开、断网也能看，**证据全部内嵌**。

    和普通 `html` 导出的区别只有一处但很关键：普通版里的出处是
    **外链**，断网或者对方站点下线之后就点不开了；单文件版把每条出处的
    标题、站点、发布时间、以及那段被引用的原文**一起写进文件**。

    🔴 **不内嵌图片**。技术上可以转 base64，但一份研究简报里内嵌
    十几张原站配图会让文件涨到几十兆，而那些图对"这句话谁说的"
    没有任何证明力 —— 内嵌的是**文字证据**，不是版面。

    🔴 **不内嵌抓来的整篇正文**。只放简报里真正引用到的那几段 ——
    整篇塞进去既是体积灾难，也把「摘录」悄悄变成了「转载」。
    """
    body = _to_html(md, title=title)
    # `_to_html` 出来是一个完整文档；这里把样式和证据区注入进去。
    # 用字符串替换而不是重新拼一个文档 —— 拼两次的话两边的排版
    # 迟早会走岔，而用户看到的是"导出的 HTML 和打印的 PDF 长得不一样"
    extra = f"<style>{_SINGLE_EXTRA_CSS}</style>"
    body = body.replace("</head>", f"{extra}</head>", 1)

    parts: list[str] = [
        '<details class="syn-raw" open><summary>原始证据（内嵌，断网也能看）</summary>'
    ]
    seen: set[str] = set()
    #: url → 它在证据区里的锚点号。给下面的 `_localize_links` 用
    anchor_by_url: dict[str, int] = {}
    for i, c in enumerate(payload.get("results") or [], 1):
        url = str(c.get("url") or "")
        if not url or url in seen:
            continue
        seen.add(url)
        anchor_by_url[url] = i
        tr = c.get("trust") or {}
        quotes = [
            str(e.get("text") or "")
            for e in _quotes_for(payload, url)
        ][:4]
        parts.append(
            f'<div class="syn-src" id="src-{i}">'
            f'<div><b>[{i}]</b> {_esc(str(c.get("title") or ""))}</div>'
            f'<div class="syn-meta">{_esc(str(c.get("site") or ""))}'
            f'　{_esc(str(c.get("published") or "") or "没有发布时间")}'
            f'　{_esc(str(tr.get("tierLabel") or "未收录"))}</div>'
            f'<div><a href="{_esc(url)}">{_esc(url)}</a></div>'
            + "".join(f"<blockquote>{_esc(q)}</blockquote>" for q in quotes)
            + "</div>"
        )
    parts.append("</details>")
    parts.append(
        '<p class="syn-meta">这份文件是自包含的：不请求任何外部资源，'
        '断网、原站下线之后仍然能看到每句话的出处和被引用的原文。'
        '**内嵌的是文字证据，不含图片和整篇正文。**</p>'
    )
    # 🔴 **这一步以前整个没有，而它是 E5/E6 的全部意义所在。**
    # 证据区有 `id="src-1"` 这样的锚点，但**没有任何东西指向它们** ——
    # 正文里用的是直接的外链 `<a href="https://…">`。所以：
    #   · 离线单文件（E6）：断网之后正文里每个来源都点不开，
    #     而那份原文其实就嵌在同一个文件的下半部分
    #   · 可点引用 PDF（E5）：链接确实被 Chromium 写进了 PDF，
    #     但它们指向**网上**，不是指向这份文件里嵌着的摘录
    # 两个功能都"能导出、排版正常、一个错都不报"，唯独承诺的那件事没做到。
    body = _localize_links(body, anchor_by_url)
    return body.replace("</body>", "".join(parts) + "</body>", 1)


def _localize_links(html: str, anchor_by_url: dict[str, int]) -> str:
    """
    把正文里指向"已经嵌进来的那些来源"的外链，改成指向文件内部的锚点。

    🔴 **只改有对应锚点的那些。** 没嵌进来的外链原样留着 ——
    把它也改成 `#src-N` 会得到一个跳不到任何地方的链接，
    而那比外链更糟：用户点了以为自己点歪了，反复点。

    🔴 **原始网址不能丢。** 改完的链接加上 `data-href`，
    并保留 `title` 提示原地址；证据区里本来就完整印着这个 URL，
    所以"这条到底来自哪"这个信息一点没少。
    """
    if not anchor_by_url:
        return html

    def repl(m: re.Match[str]) -> str:
        url = m[2]
        n = anchor_by_url.get(url)
        if n is None:
            return m[0]
        # 去掉 target=_blank：跳的是本文件内部，开新窗口毫无意义
        return f'<a class="syn-cite" href="#src-{n}" data-href="{url}" title="{url}">'

    return re.sub(r'(<a\b[^>]*?\bhref=")([^"]+)("[^>]*>)', repl, html, flags=re.I)


def _quotes_for(payload: dict[str, Any], url: str) -> list[dict[str, Any]]:
    """从简报里把引用了这条 url 的摘录挑出来。"""
    out: list[dict[str, Any]] = []

    def walk(node: Any) -> None:
        if isinstance(node, dict):
            if str(node.get("url") or "") == url and node.get("text"):
                out.append(node)
            for v in node.values():
                walk(v)
        elif isinstance(node, list):
            for v in node:
                walk(v)

    walk(payload.get("briefing") or {})
    return out
